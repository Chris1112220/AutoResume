from flask import Flask, jsonify, send_file
from config import SQLALCHEMY_DATABASE_URI
from models import db, Education, Company, Job, Accomplishment, TechnicalSkill, Project
from flask import request, render_template

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = SQLALCHEMY_DATABASE_URI
db.init_app(app)


def extract_keywords(text):
    stopwords = {'and', 'the', 'in', 'on', 'with', 'an', 'a',
                 'to', 'for', 'we', 'are', 'of', 'is', 'be', 'as'}
    words = text.lower().split()
    keywords = [word.strip(".,")
                for word in words if word not in stopwords and len(word) > 2]
    return set(keywords)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_tab_stop(paragraph, position_twips):
    pPr = paragraph._p.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs')
        pPr.append(tabs)
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(position_twips))
    tabs.append(tab)


@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')


@app.route('/resume', methods=['POST'])
def generate_resume():
    job_description = """
    Seeking an RPA Developer with expertise in UiPath, automation, and strong problem-solving skills.
    """
    keywords = extract_keywords(job_description)

    matched_accomplishments = []
    all_accomplishments = Accomplishment.query.all()

    tech_skills = TechnicalSkill.query.all()
    projects = Project.query.all()
    jobs = db.session.query(Job, Company).join(
        Company, Job.company_id == Company.id).all()

    doc = Document()

    # HEADER
    doc.add_heading("Christopher A. Roberts", level=0)
    header_para = doc.add_paragraph()
    header_para.paragraph_format.space_after = Pt(0)
    header_para.paragraph_format.space_before = Pt(0)
    header_para.add_run("Philadelphia, PA | ").font.name = 'Calibri'
    header_para.add_run(
        "Christopher.roberts11220@gmail.com | ").font.name = 'Calibri'
    header_para.add_run("908-963-0613 | ").font.name = 'Calibri'
    add_hyperlink(header_para, "https://github.com/Chris1112220", "GitHub")
    header_para.add_run(" | ").font.name = 'Calibri'
    add_hyperlink(
        header_para, "https://www.linkedin.com/in/christopher-roberts-philadelphia/", "LinkedIn")

    # TECHNICAL SKILLS AND AWARDS
    doc.add_heading("Technical Skills and Awards", level=1)
    tech_para = doc.add_paragraph()
    tech_para.paragraph_format.space_after = Pt(0)

    tech_para.add_run("Programming Languages: ").bold = True
    tech_para.add_run("Python, Java, C\n")
    tech_para.add_run("Automation Tools: ").bold = True
    tech_para.add_run("UiPath Studio, StudioX, Orchestrator\n")
    tech_para.add_run("Databases: ").bold = True
    tech_para.add_run("MySQL, PostgreSQL, SQLite\n")
    tech_para.add_run("Cloud Platforms: ").bold = True
    tech_para.add_run("AWS, Azure\n")
    tech_para.add_run("Version Control: ").bold = True
    tech_para.add_run("GitHub\n")
    tech_para.add_run("Awards: ").bold = True
    tech_para.add_run("Streamline Employee of the Quarter 2019")

    # PROFESSIONAL EXPERIENCE
    doc.add_heading("Professional Experience", level=1)

    company_dates = {
        "Drexel University": "March 2021 – Present",
        "Columbus Construction": "August 2020 – December 2020",
        "Streamline": "April 2018 – May 2020",
        "CTI Foods": "November 2014 – March 2018"
    }

    for job, company in jobs:
        job_header = doc.add_paragraph()
        job_header.paragraph_format.space_after = Pt(0)
        p_run = job_header.add_run(f"{company.name}, {company.location}")
        p_run.bold = True
        add_tab_stop(job_header, 9360)
        job_header.add_run("\t")
        date_run = job_header.add_run(company_dates.get(company.name, ""))

        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_after = Pt(0)
        title_run = title_para.add_run(job.title)
        title_run.italic = True

        # Always show accomplishments tied to the job
        job_accomplishments = [
            acc for acc in all_accomplishments if acc.job_id == job.id]

        for acc in job_accomplishments:
            bullet = doc.add_paragraph(acc.description, style='List Bullet')
            bullet.paragraph_format.space_after = Pt(0)

        doc.add_paragraph()  # Add space between professional experiences

        # PROJECTS
    projects_para = doc.add_paragraph()
    projects_run = projects_para.add_run("Projects | ")
    projects_run.bold = True
    add_hyperlink(
        projects_para, "https://chris-dev-portfolio-one.vercel.app/", "Portfolio")

    # List projects separately (bold project names)
    for proj in projects:
        project_para = doc.add_paragraph()
        name_run = project_para.add_run(proj.name + " – ")
        name_run.bold = True
        desc_run = project_para.add_run(proj.description)
        desc_run.font.name = 'Calibri'
        project_para.paragraph_format.space_after = Pt(0)

    # EDUCATION
    doc.add_heading("Education", level=1)

    drexel_para = doc.add_paragraph()
    drexel_para.paragraph_format.space_after = Pt(6)
    drexel_para.paragraph_format.space_before = Pt(0)
    add_tab_stop(drexel_para, 9360)
    drexel_run1 = drexel_para.add_run(
        "Drexel University, College of Computing and Informatics, Philadelphia, PA")
    drexel_run1.bold = True
    drexel_para.add_run("\tJanuary 2024")
    drexel_para.add_run(
        "\nPost-Baccalaureate Graduate Certificate in Computer Science Foundations")

    temple_para = doc.add_paragraph()
    temple_para.paragraph_format.space_after = Pt(0)
    temple_para.paragraph_format.space_before = Pt(6)
    add_tab_stop(temple_para, 9360)
    temple_run1 = temple_para.add_run(
        "Temple University, Fox School of Business — Philadelphia, PA")
    temple_run1.bold = True
    temple_para.add_run("\tJanuary 2012")
    temple_para.add_run("\nBBA, Finance")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="Christopher_Roberts_Resume.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == '__main__':
    app.run(debug=True)
